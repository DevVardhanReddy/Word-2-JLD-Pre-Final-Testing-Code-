from io import BytesIO
import pandas as pd
import docx
import os
from openpyxl import Workbook
import glob
import sys

from PyQt5.QtWidgets import QApplication, QWidget, QVBoxLayout, QLabel, QPushButton, QFileDialog, QLineEdit, QHBoxLayout, QGraphicsDropShadowEffect

from PyQt5.QtGui import QPixmap, QColor

from PyQt5.QtCore import QThread, Qt, pyqtSignal

from PyQt5.QtGui import QFont

from PyQt5.QtWidgets import QWidget, QLabel, QPushButton, QVBoxLayout, QHBoxLayout, QLineEdit, QFileDialog, QSizePolicy, QSpacerItem

from PyQt5.QtWidgets import QComboBox

from PyQt5.QtWidgets import (

    QApplication, QWidget, QVBoxLayout, QHBoxLayout, QLabel, QPushButton,

    QComboBox, QLineEdit, QFileDialog, QSpacerItem, QSizePolicy

)

from PyQt5.QtGui import QPixmap

from PyQt5.QtCore import Qt

import sys

import openpyxl

import pandas as pd

import os

import xml.etree.ElementTree as ET

from openpyxl import Workbook

from openpyxl import load_workbook

import os

import docx

from docx import Document

from openpyxl import Workbook

import sys

import os

import json

import pandas as pd

import xml.etree.ElementTree as ET

from PyQt5.QtWidgets import (QApplication, QWidget, QVBoxLayout, QLabel, QPushButton, QFileDialog,



                            QLineEdit, QHBoxLayout, QGraphicsDropShadowEffect, QComboBox,



                            QSizePolicy, QSpacerItem, QTextEdit, QProgressBar, QFrame)

from PyQt5.QtGui import QPixmap, QColor, QFont

from PyQt5.QtCore import Qt, QTimer

from functools import partial

import json

import warnings

from flask import Flask, request, jsonify, render_template

from flask_cors import CORS

import os









class StyleExtractor:

   

    def extract_data(self, input_path):



        if not input_path:



            return False, "Error: No valid file or folder selected!"



        # Determine if the input is a file or a folder



        if os.path.isfile(input_path) and input_path.endswith(".docx"):



            # Single file case



            folder_path = os.path.dirname(input_path)



            output_excel = os.path.join(folder_path, "Excel_Output", "StylesData.xlsx")



            success = self.process_documents([input_path], output_excel)



        elif os.path.isdir(input_path):



            # Folder case - Process all .docx files inside



            output_excel = os.path.join(input_path, "Excel_Output", "StylesData.xlsx")



            doc_files = [

                os.path.join(input_path, f)

                for f in os.listdir(input_path)

                if f.endswith(".docx")

            ]



            success = self.process_documents(doc_files, output_excel)



        else:



            return False, "Error: Invalid file or folder path!"



        if success:



            return True, "Extraction completed successfully!"



        else:



            return False, "Extraction failed."



    def process_documents(self, doc_files, output_excel):



        try:



            wb = Workbook()



            ws = wb.active



            ws.title = "Styled Text"



            ws.append(

                [

                    "Document Name",

                    "Paragraph Number",

                    "Paragraph Style",

                    "Text Style",

                    "Text",

                ]

            )



            for doc_path in doc_files:



                doc_name = os.path.splitext(os.path.basename(doc_path))[0]



                try:



                    doc = docx.Document(doc_path)



                except Exception as e:



                    print(f"Error opening {doc_name}: {e}")



                    continue  # Skip to the next document



                styled_content = self.extract_text_with_styles(doc)



                for content in styled_content:



                    paragraph_num, para_style, text_style, text = content



                    ws.append([doc_name, paragraph_num, para_style, text_style, text])



            # Ensure output folder exists



            output_folder = os.path.dirname(output_excel)



            os.makedirs(output_folder, exist_ok=True)



            wb.save(output_excel)



            return True



        except Exception as e:



            print(f"Error processing documents: {e}")



            return False



    def get_colour(self, run):



        colour = run.font.color



        if colour and colour.rgb:



            rgb = (colour.rgb[0], colour.rgb[1], colour.rgb[2])



        else:



            return ""



        colour_dictionary = {

            (255, 0, 0): "Red",

            (0, 255, 0): "Green",

            (0, 0, 255): "Blue",

            (255, 255, 0): "Yellow",

            (128, 128, 128): "Gray",

            (255, 153, 0): "Orange",

            (0, 0, 0): "",

        }



        colour_value = colour_dictionary.get(rgb, rgb)



        return colour_value



    def get_paragraph_style(self, paragraph):



        alignment_map = {0: "Left", 1: "Centre", 2: "Right", 3: "Justify"}



        alignment = alignment_map.get(paragraph.alignment, "Left")



        styles = [

            f"_LeftInd:{paragraph.paragraph_format.left_indent.pt}pt"

            if paragraph.paragraph_format.left_indent

            else "",

            f"_RightInd:{paragraph.paragraph_format.right_indent.pt}pt"

            if paragraph.paragraph_format.right_indent

            else "",

            f"_SpaceBefore:{paragraph.paragraph_format.space_before.pt}pt"

            if paragraph.paragraph_format.space_before

            else "",

            f"_SpaceAfter:{paragraph.paragraph_format.space_after.pt}pt"

            if paragraph.paragraph_format.space_after

            else "",

            f"_LineSpacing:{paragraph.paragraph_format.line_spacing}pt"

            if paragraph.paragraph_format.line_spacing

            else "",

            f"_Hanging:{paragraph.paragraph_format.first_line_indent.pt}pt"

            if paragraph.paragraph_format.first_line_indent

            else "",

        ]



        return alignment + "".join(filter(None, styles))



    def get_text_style(self, run):



        font_name = run.font.name or "Arial"



        font_size = f"_{run.font.size.pt}pt" if run.font.size else "_12.0pt"



        bold = "_Bold" if run.bold else ""



        italic = "_Italic" if run.italic else ""



        underline = "_Underline" if run.underline else ""



        # Get color using the class method



        colour_value = self.get_colour(run)



        colour = f"_{colour_value}" if colour_value else ""



        return font_name + font_size + bold + italic + underline + colour



    def extract_text_with_styles(self, doc):



        content = []



        last_para_style = None



        last_text_style = None



        for i, paragraph in enumerate(doc.paragraphs):



            para_style = self.get_paragraph_style(paragraph) or last_para_style



            paragraph_num = i + 1



            if not paragraph.text.strip():  # Empty paragraph



                content.append((paragraph_num, para_style, last_text_style, ""))



                continue



            runs = paragraph.runs



            current_text = ""



            current_style = None



            for run in runs:



                text = run.text



                text_style = self.get_text_style(run) or last_text_style



                if current_style == text_style:



                    current_text += text



                else:



                    if current_text:



                        content.append(

                            (paragraph_num, para_style, current_style, current_text)

                        )



                    current_text = text



                    current_style = text_style



            if current_text:



                content.append((paragraph_num, para_style, current_style, current_text))



            last_para_style = para_style



            last_text_style = current_style



        return content





class WordToJLDConverter:
    def __init__(self):
        self.style_extractor = StyleExtractor()

    def process_files(self, word_files, output_path):
        try:
            # Step 1: Convert Word to Excel in memory
            excel_data = self._convert_word_to_excel(word_files)
            
            # Step 2: Convert Excel to JLD
            jld_data = self._convert_excel_to_jld(excel_data)
            
            # Step 3: Save JLD output
            return self._save_jld(jld_data, output_path)
            
        except Exception as e:
            return False, str(e)

    def _convert_word_to_excel(self, word_files):
        wb = Workbook()
        ws = wb.active
        ws.append(["Document", "Paragraph", "Style", "Text"])
        
        for file in word_files:
            doc = docx.Document(BytesIO(file.read()))
            content = self.style_extractor.extract_text_with_styles(doc)
            
            for para_num, style, text_style, text in content:
                ws.append([file.filename, para_num, style, text_style, text])

        excel_data = BytesIO()
        wb.save(excel_data)
        excel_data.seek(0)
        return excel_data

    def _convert_excel_to_jld(self, excel_data):
        # Implement your actual Excel-to-JLD conversion here
        jld_data = BytesIO()
        jld_data.write(b"JLD_FORMATTED_OUTPUT")
        jld_data.seek(0)
        return jld_data

    def _save_jld(self, jld_data, output_path):
        os.makedirs(output_path, exist_ok=True)
        output_file = os.path.join(output_path, "output.jld")
        
        with open(output_file, 'wb') as f:
            f.write(jld_data.read())
            
        return True, output_file



COLOR_NAME_MAP = {    (255, 0, 0): "Red",    (0, 255, 0): "Green",    (0, 0, 255): "Blue",    (128, 128, 128): "Gray",    (255,255,0): "Yellow",    (255,0,255): "Magenta",    (0,255,255): "Cyan",    (128,0,0): "Dark Red",    (0,128,0): "Dark Green",    (0,0,128): "Dark Blue",    (128,128,0): "Olive",    (128,0,128): "Purple",    (0,128,128): "Teal",    (192,192,192): "Silver",    (128,128,128): "Gray",    (255,255,224): "Light Yellow",    (144,238,144): "Light Green",    (224,255,255): "Light Turquoise",    (173,216,230): "Light Blue",    (230,230,250): "Light Purple",    (255,182,193): "Light Red",    (255,240,245): "Light Magenta",    (240,255,255): "Light Cyan",    (211,211,211): "Light Gray",    (192,192,128): "Light Olive",    (128,192,192): "Light Teal",    (224,224,224): "Light Silver",    (64,64,64): "Light Black",    (245,245,245): "Light White",    (210,180,140): "Light Brown",    (255,165,0): "Light Orange",    (255,192,203): "Light Pink",    (255,215,0): "Light Gold",    (205,127,50): "Light Bronze",    (184,115,51): "Light Copper",    (181,166,66): "Light Brass",    (192,192,192): "Light Silver",    (169,169,169): "Light Gray",    (119,136,153): "Light Slate",    (176,196,222): "Light Steel",    (135,206,235): "Light Sky",    (32,178,170): "Light Sea",    (0,255,127): "Light Spring",    (189,252,201): "Light Mint",    (50,205,50): "Light Lime",    (34,139,34): "Light Forest",    (107,142,35): "Light Olive",    (240,230,140): "Light Khaki",    (245,245,220): "Light Beige",    (245,222,179): "Light Wheat",    (210,180,140): "Light Tan",    (255,218,185): "Light Peach",    (240,128,128): "Light Coral",    (250,128,114): "Light Salmon",    (255,99,71): "Light Tomato",}







def extract_colors_from_docx(docx_path):

    """

    Extract unique colors from a DOCX file.

    Returns a set of integer RGB tuples (0-255), skipping black (0,0,0).

    """

    colors = set()



    doc = Document(docx_path)



    for para in doc.paragraphs:



        for run in para.runs:



            if run.font.color and run.font.color.rgb:



                hex_value = str(run.font.color.rgb)  # e.g., "FF0000"



                try:



                    r = int(hex_value[0:2], 16)



                    g = int(hex_value[2:4], 16)



                    b = int(hex_value[4:6], 16)



                    if (r, g, b) == (0, 0, 0):  # skip black since it's default



                        continue



                    colors.add((r, g, b))



                except Exception as e:



                    print(f"Error converting color in {docx_path}: {e}")



    return colors



def extract_fonts_from_docx(docx_path):



    """



    Extract unique font names from a DOCX file.



    Returns a set of font names.



    """



    fonts = set()



    doc = Document(docx_path)



    for para in doc.paragraphs:



        for run in para.runs:



            if run.font and run.font.name:



                fonts.add(run.font.name)



    return fonts



def extract_parastyles_from_docx(docx_path):



    """



    Extract unique paragraph styles based on formatting from a DOCX file.



    For each paragraph, we capture:



      - Alignment (use the string value; default "Left" if not set)



      - LeftIndent (in points; default 0)



      - RightIndent (default 0)



      - SpaceBefore (default 0)



      - SpaceAfter (default 0)



      - LineSpacing (default 0)







    Returns a set of tuples:



        (alignment, left_indent, right_indent, space_before, space_after, line_spacing)



    """



    styles = set()



    doc = Document(docx_path)



    for para in doc.paragraphs:



        p_format = para.paragraph_format



        # Use string values; if not set, assign defaults.



        alignment = p_format.alignment



        # Map docx.enum.text.WD_ALIGN_PARAGRAPH values to strings.



        if alignment is None:



            alignment_str = "Left"



        else:



            # Use the built-in name if available.



            alignment_str = alignment.name.capitalize() if hasattr(alignment, 'name') else str(alignment)



        left_indent = p_format.left_indent.pt if p_format.left_indent is not None else 0



        right_indent = p_format.right_indent.pt if p_format.right_indent is not None else 0



        space_before = p_format.space_before.pt if p_format.space_before is not None else 0



        space_after = p_format.space_after.pt if p_format.space_after is not None else 0



        line_spacing = p_format.line_spacing if p_format.line_spacing is not None else 0



        # Create a tuple (rounded to 2 decimals for consistency)



        style_tuple = (



            alignment_str,



            round(left_indent, 2),



            round(right_indent, 2),



            round(space_before, 2),



            round(space_after, 2),



            round(line_spacing, 2) if isinstance(line_spacing, (int, float)) else 0



        )



        styles.add(style_tuple)



    return styles



def generate_color_xml_tag(color_tuple):



    """



    Generates an XML <Color> element from an integer RGB tuple.



    """



    color_name = COLOR_NAME_MAP.get(



        color_tuple,



        f"Color_{color_tuple[0]}_{color_tuple[1]}_{color_tuple[2]}"



    )



    color_id = color_name  # Use friendly name as identifier.







    color_elem = ET.Element("Color")



    id_elem = ET.SubElement(color_elem, "Id", Name=color_id)



    id_elem.text = f"Def.{color_id}"







    name_elem = ET.SubElement(color_elem, "Name")



    name_elem.text = color_name







    # Use normalized RGB values (0 to 1) in the XML.



    normalized_rgb = (color_tuple[0] / 255.0, color_tuple[1] / 255.0, color_tuple[2] / 255.0)



    rgb_elem = ET.SubElement(color_elem, "RGB")



    rgb_elem.text = f"{normalized_rgb[0]},{normalized_rgb[1]},{normalized_rgb[2]}"







    # Add additional sub-tags with dummy values.



    ET.SubElement(color_elem, "CMYK").text = "0,0,0,0"



    ET.SubElement(color_elem, "LAB").text = "0,0,0"



    ET.SubElement(color_elem, "HSB").text = "0,0,0"



    ET.SubElement(color_elem, "SpotColor").text = "0"



    ET.SubElement(color_elem, "OverwriteSpotColor").text = "True"



    ET.SubElement(color_elem, "SeparationColorSpace")



    ET.SubElement(color_elem, "Density").text = "100"



    ET.SubElement(color_elem, "IsDeviceN").text = "False"



    ET.SubElement(color_elem, "IsInherit").text = "False"



    ET.SubElement(color_elem, "ColorType").text = "Simple"







    return color_elem



def generate_fill_xml_tag(color_tuple):



    """



    Generates an XML <FillStyle> element corresponding to the given color,



    with sub-elements in the order: ID, Name, ColorId.



    """



    color_name = COLOR_NAME_MAP.get(



        color_tuple,



        f"Color_{color_tuple[0]}_{color_tuple[1]}_{color_tuple[2]}"



    )



    fill_name = f"{color_name}Fill"







    fill_elem = ET.Element("FillStyle")



    id_elem = ET.SubElement(fill_elem, "Id", Name=fill_name)



    id_elem.text = f"Def.{fill_name}"







    name_elem = ET.SubElement(fill_elem, "Name")



    name_elem.text = fill_name







    colorid_elem = ET.SubElement(fill_elem, "ColorId")



    colorid_elem.text = f"Def.{color_name}"







    return fill_elem



def generate_font_xml_tag(font_name):



    """



    Generates an XML <Font> element for the given font name.



    """



    font_elem = ET.Element("Font")



    id_elem = ET.SubElement(font_elem, "Id", Name=font_name)



    id_elem.text = f"Def.{font_name}"







    name_elem = ET.SubElement(font_elem, "Name")



    name_elem.text = font_name







    fontname_elem = ET.SubElement(font_elem, "FontName")



    fontname_elem.text = font_name







    subfont_elem = ET.SubElement(font_elem, "SubFont", Name="Regular", Bold="False", Italic="False")



    ET.SubElement(subfont_elem, "FontIndex").text = "0"



    ET.SubElement(subfont_elem, "FontLocation").text = f"FONT_DIR,{font_name}.ttf"



    ET.SubElement(subfont_elem, "OpenFontFileFlag").text = "68"







    ET.SubElement(font_elem, "IsSymbolFont").text = "False"







    return font_elem



def generate_parastyle_xml_tag(style_tuple):



    """



    Generates an XML <ParaStyle> element based on the style_tuple:



      (alignment, left_indent, right_indent, space_before, space_after, line_spacing)



    The style name is generated in the format:



      Alignment_LeftIndent_RightIndent_SpaceBefore_SpaceAfter_LineSpacing



    For example: Center_0_0_0_0_0



    The XML element is built as:







    <ParaStyle>



        <Id Name="Center_0_0_0_0_0">Def.Center_0_0_0_0_0ParaStyle</Id>



        <Name>Center_0_0_0_0_0</Name>



        <LeftIndent>0</LeftIndent>



        <RightIndent>0</RightIndent>



        <SpaceBefore>0</SpaceBefore>



        <SpaceAfter>0</SpaceAfter>



        <LineSpacing>0</LineSpacing>



        <HAlign>Center</HAlign>



    </ParaStyle>



    """



    alignment, left_indent, right_indent, space_before, space_after, line_spacing = style_tuple



    style_name = f"{alignment}_{left_indent}_{right_indent}_{space_before}_{space_after}_{line_spacing}"







    para_elem = ET.Element("ParaStyle")



    id_elem = ET.SubElement(para_elem, "Id", Name=style_name)



    id_elem.text = f"Def.{style_name}ParaStyle"







    name_elem = ET.SubElement(para_elem, "Name")



    name_elem.text = style_name







    ET.SubElement(para_elem, "LeftIndent").text = str(left_indent)



    ET.SubElement(para_elem, "RightIndent").text = str(right_indent)



    ET.SubElement(para_elem, "SpaceBefore").text = str(space_before)



    ET.SubElement(para_elem, "SpaceAfter").text = str(space_after)



    ET.SubElement(para_elem, "LineSpacing").text = str(line_spacing)



    if alignment == "Justify":



        alignment = "JustifyLeft"



    ET.SubElement(para_elem, "HAlign").text = alignment







    return para_elem



def process_documents_and_update_xml(docx_folder, xml_input_path, xml_output_path):



    # 1. Extract colors, fonts, and paragraph styles from all DOCX files.



    docx_files = glob.glob(os.path.join(docx_folder, "*.docx"))



    all_colors = set()



    all_fonts = set()



    all_parastyles = set()







    for docx_file in docx_files:



        all_colors.update(extract_colors_from_docx(docx_file))



        all_fonts.update(extract_fonts_from_docx(docx_file))



        all_parastyles.update(extract_parastyles_from_docx(docx_file))







    # 2. Generate XML elements for new colors and corresponding fill styles.



    new_color_and_fill = []



    for color in all_colors:



        color_elem = generate_color_xml_tag(color)



        fill_elem = generate_fill_xml_tag(color)



        new_color_and_fill.append((color_elem, fill_elem))







    # 3. Generate XML elements for new fonts.



    new_font_elements = [generate_font_xml_tag(font) for font in all_fonts]







    # 4. Generate XML elements for new paragraph styles.



    new_para_elements = [generate_parastyle_xml_tag(style) for style in all_parastyles]







    # 5. Parse the original XML file.



    tree = ET.parse(xml_input_path)



    root = tree.getroot()







    # 6. Determine the insertion point (we insert before <MessageManagement>).



    insert_index = None



    for idx, child in enumerate(root):



        if child.tag == "MessageManagement":



            insert_index = idx



            break



    if insert_index is None:



        insert_index = len(root)







    # 7. Insert new Color elements and immediately after each, its FillStyle element.



    existing_color_ids = {c.find("Id").attrib.get("Name") for c in root.findall("Color") if c.find("Id") is not None}



    for color_elem, fill_elem in new_color_and_fill:



        new_id = color_elem.find("Id").attrib.get("Name")



        if new_id not in existing_color_ids:



            root.insert(insert_index, color_elem)



            insert_index += 1



            root.insert(insert_index, fill_elem)



            insert_index += 1







    # 8. Insert new Font elements.



    existing_font_ids = {f.find("Id").attrib.get("Name") for f in root.findall("Font") if f.find("Id") is not None}



    for font_elem in new_font_elements:



        new_id = font_elem.find("Id").attrib.get("Name")



        if new_id not in existing_font_ids:



            root.insert(insert_index, font_elem)



            insert_index += 1







    # 9. Insert new ParaStyle elements after the Font elements.



    existing_para_ids = {p.find("Id").attrib.get("Name") for p in root.findall("ParaStyle") if p.find("Id") is not None}



    for para_elem in new_para_elements:



        new_id = para_elem.find("Id").attrib.get("Name")



        if new_id not in existing_para_ids:



            root.insert(insert_index, para_elem)



            insert_index += 1





    # 10. Write the updated XML to file.

    tree.write(xml_output_path, encoding="utf-8", xml_declaration=True)



    print(



        f"Inserted {len(new_color_and_fill)} new colors (and fill styles), {len(new_font_elements)} fonts, and {len(new_para_elements)} paragraph style(s) into {xml_output_path}")





class WordToJLDConverter:



    def __init__(self):

        # Initialize paths

        self.baseblock_path = r"c:\Users\Desktop\Word-JLD Original File\BaseBlockHK1.txt"

        self.variables_file_path = r"c:\UsersDesktop\Word-JLD Original File\VariablesList.xlsx"

        self.cancel_requested = False



    def read_excel(self, file_path):

        wb = openpyxl.load_workbook(file_path)

        ws = wb.active



        document_data = {}



        # read rows

        for row in ws.iter_rows(min_row=2, values_only=True):

            document_name, paragraph_number, paragraph_style, text_style, text = row



            if document_name not in document_data:

                document_data[document_name] = []



            document_data[document_name].append({

                "paragraph_number": paragraph_number,

                "paragraph_style": paragraph_style,

                "text_style": text_style,

                "text": text

            })



        return document_data



    def get_var_update(self, variables_list_path):

        workbook = load_workbook(variables_list_path, data_only=True)

        sheet = workbook.active

        variables_array = []



        for row in sheet.iter_rows(min_row=2):

            cell_value_B = row[1].value

            cell_value_F = row[5].value



            if cell_value_B is not None:

                variables_array.append(str(cell_value_B))

            if cell_value_F is not None:

                variables_array.append(str(cell_value_F))



        result_array = ",\n".join(variables_array)



        return result_array



    def get_variables_list(self, variables_list_path):

        workbook = load_workbook(variables_list_path)

        sheet = workbook.active

        variables_array = []



        for row in sheet.iter_rows(min_row=2):

            cell_value_A = row[0].value

            cell_value_E = row[4].value



            formatted_A = f"{{Form.{cell_value_A}}}" if cell_value_A is not None else None

            formatted_E = f"{{{cell_value_E}}}" if cell_value_E is not None else None



            if formatted_A:

                variables_array.append(formatted_A)

            if formatted_E:

                variables_array.append(formatted_E)



        return variables_array



    def create_section_flow(self, section_content, count):

        prefix1 = '{"LocalPath": ["Flows","Flow '

        prefix2 = '"],"Class": "Flow","LockedWebNode": true,"Type": "Simple","Width": 0.19,"Content": ['

        suffix = '],"DocxLock": "Inherit","IsInsertPoint": false,"SectionFlow": false,"FlowUsageLogging": false,"AllowBlockOnlyOnce": false,"LowercaseFirstCharacterOfNextFlow": false,"AcceptLowercasingRuleFromPreviousFlow": false}'



        new_section_content = [f'"{item}"' if isinstance(item, str) else str(item) for item in section_content]

        combined = ','.join(new_section_content)



        section_flow = prefix1 + str(count) + prefix2 + str(combined) + suffix

        return section_flow



    def process_file_or_folder(self, path):

        """Processes a single file or all files in a folder."""

        # Ensure required paths are set

        if not hasattr(self, 'output_folder') or not self.output_folder:

            self.output_folder = os.path.join(path, "JLD_Output")



        if os.path.isfile(path):

            # Process single file

            self.process_excel_file(path)

        elif os.path.isdir(path):

            # Process all files in the folder

            excel_files = self.get_all_excel_files(path)



            if not excel_files:

                print("❌ No Excel files found in the selected folder.")

                return



            for file in excel_files:

                print(f"Selected file: {file}")

                try:

                    document_data = self.read_excel(file)

                    self.create_jld_files(document_data, self.baseblock_path, self.output_folder, self.variables_file_path)

                    print(f"✅ Processing complete for: {file}")

                except Exception as e:

                    print(f"❌ Error processing file {file}: {e}")



            print("✅ Processing complete for all Excel files in the folder.")



    def process_excel_file(self, file_path):

        """Reads and processes a single Excel file"""

        try:

            document_data = self.read_excel(file_path)

            self.create_jld_files(document_data, self.baseblock_path, self.output_folder, self.variables_file_path)

            print(f"✅ Processing complete for: {file_path}")

        except Exception as e:

            print(f"❌ Error processing file {file_path}: {e}")



    def create_jld_files(self, document_data, baseblock_path, output_folder, variables_file_path):

        if not os.path.exists(output_folder):

            os.makedirs(output_folder)



        variables_list = self.get_variables_list(variables_file_path)



        # read base block

        with open(baseblock_path, 'r') as file:

            baseblock_template = file.read()



        for document_name, paragraphs_info in document_data.items():

            unique_paragraph_styles = []

            unique_text_styles = []

            content_list = []



            section_flow_count = 3

            for para_info in paragraphs_info:

                text_val = para_info["text"]

                if text_val and text_val.startswith("{NewFlow}"):

                    section_flow_count += 1



            flow_indices = list(range(1, section_flow_count + 1))



            # unique styles

            for para_info in paragraphs_info:

                para_style = para_info["paragraph_style"]

                text_style = para_info["text_style"]



                if para_style not in unique_paragraph_styles:

                    unique_paragraph_styles.append(para_style)

                if text_style not in unique_text_styles:

                    unique_text_styles.append(text_style)



            # indexing

            paragraph_index_map = {style: idx + 1 for idx, style in enumerate(unique_paragraph_styles)}

            text_index_start = len(unique_paragraph_styles) + 1

            text_index_map = {style: idx + text_index_start for idx, style in enumerate(unique_text_styles)}



            section_content = []

            section_flow = ""

            new_section_flow = ""

            count = 1



            last_paragraph_number = None



            for para_info in paragraphs_info:

                para_style = para_info["paragraph_style"]

                text_style = para_info["text_style"]

                text = para_info["text"]

                if text is None:

                    text = ""



                for index, vari in enumerate(variables_list, start=1):

                    if vari in text:

                        text = text.replace(vari, "\"," + str(index + len(unique_paragraph_styles) + len(unique_text_styles) + len(flow_indices)) + ",\"")



                para_style_idx = paragraph_index_map[para_style] + len(flow_indices)

                text_style_idx = text_index_map[text_style] + len(flow_indices)



                current_paragraph_number = para_info["paragraph_number"]

                if current_paragraph_number != last_paragraph_number:

                    content_list.extend([para_style_idx])

                last_paragraph_number = current_paragraph_number

                if "{NewFlow}" not in text:

                    content_list.extend([text_style_idx, text])

                else:

                    section_flow = self.create_section_flow(content_list, count)

                    count += 1

                    new_section_flow += '\n' + section_flow + ','

                    first_text = text[9:].strip()

                    content_list = [para_style_idx, text_style_idx, first_text]



            section_flow = self.create_section_flow(content_list, count)

            new_section_flow += '\n' + section_flow + ','



            para_styles_string = ', '.join(

                [f'{{"NodePath": ["ParagraphStyles", "{style}"], "Cls": "ParaStyle"}}' for style in unique_paragraph_styles])

            text_styles_string = ', '.join(

                [f'{{"NodePath": ["TextStyles", "{style}"], "Cls": "TextStyle"}}' for style in unique_text_styles])



            fidx_str = ",".join(map(str, flow_indices))



            var_update = self.get_var_update(variables_file_path)


        modified_baseblock = baseblock_template.replace('PARAUPDATE', para_styles_string) \
                .replace('TEXTUPDATE', text_styles_string) \
                .replace('CONTENTVAL', new_section_flow) \
                .replace('SECTIONFLOWCOUNT', fidx_str) \
                .replace('VARUPDATE', var_update)

        output_filename = f"{document_name}.jld"
        output_path = os.path.join(output_folder, output_filename)



        with open(output_path, 'w') as output_file:

                output_file.write(modified_baseblock)



        print(f"✅ Conversion successful: {output_path}")



    def get_all_excel_files(self, folder_path):

        """Finds all Excel files in a given folder"""

        return [os.path.join(folder_path, f) for f in os.listdir(folder_path) if f.endswith(".xlsx")]



    def convert_excel_to_jld(self, input_file, output_file):

        """

        Converts an Excel (.xlsx) file into a JSON-based JLD format.

       

        :param input_file: Path to the input Excel file.

        :param output_file: Path to save the generated JLD file.

        :return: Tuple (success: bool, message: str)

        """

        try:

            # Load the Excel file

            wb = openpyxl.load_workbook(input_file)

            sheet = wb.active  # Get the first sheet



            content = []

            for row in sheet.iter_rows(values_only=True):

                if any(row):  # Ensures empty rows are skipped

                    content.append({

                        "type": "row",

                        "content": [str(cell) if cell is not None else "" for cell in row]

                    })



            # Construct the JSON structure for JLD output

            jld_data = {

                "document_name": os.path.splitext(os.path.basename(input_file))[0],  # Remove .xlsx extension

                "content": content

            }



            # Save as JSON

            with open(output_file, "w", encoding="utf-8") as json_file:

                json.dump(jld_data, json_file, indent=4)



            return True, f"✅ Conversion successful: {output_file}"



        except Exception as e:

            return False, f"❌ Error: {str(e)}"



# Example of how to use the converter

def main():

    warnings.simplefilter(action='ignore', category=UserWarning)

   

    converter = WordToJLDConverter()

   

    # Set the output folder

    converter.output_folder = r"C:\path\to\output\folder\JLD_Output"

   

    # Process a single file

    file_path = r"C:\path\to\input\file.xlsx"

    converter.process_excel_file(file_path)

   

    # Or process a folder

    # folder_path = r"C:\path\to\folder"

    # converter.process_file_or_folder(folder_path)