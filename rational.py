from flask import Blueprint, request, flash, redirect, url_for, session, render_template, send_from_directory
from werkzeug.utils import secure_filename
import os
import docx
import pandas as pd
import numpy as np
from openpyxl import Workbook
from openpyxl.styles import Font
from openpyxl.worksheet.table import Table, TableStyleInfo
from datetime import datetime
from sklearn.metrics.pairwise import cosine_similarity
import cohere
from flask import current_app

# Create blueprint
web_bp = Blueprint('web', __name__)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in current_app.config['ALLOWED_EXTENSIONS']

def read_word_doc(file_path):
    doc = docx.Document(file_path)
    return "\n".join(para.text for para in doc.paragraphs)

def process_uploaded_documents():
    docs = []
    doc_names = []
    
    for filename in sorted(os.listdir(current_app.config['UPLOAD_FOLDER'])):
        if filename.endswith('.docx'):
            file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
            docs.append(read_word_doc(file_path))
            doc_names.append(filename)
    
    return docs, doc_names

def save_similarity_report(similarity_matrix, doc_names):
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    report_filename = f'similarity_report_{timestamp}.xlsx'
    report_path = os.path.join(current_app.config['REPORT_FOLDER'], report_filename)

    df = pd.DataFrame(similarity_matrix, index=doc_names, columns=doc_names)
    df.index.name = 'Document Names'

    with pd.ExcelWriter(report_path, engine='openpyxl') as writer:
        df.to_excel(writer, sheet_name='Report')
        workbook = writer.book
        sheet = workbook['Report']

        # Formatting
        bold_font = Font(bold=True)
        for cell in sheet[1]:
            cell.font = bold_font
        sheet['A1'].font = bold_font

        # Convert to percentages
        for row in sheet.iter_rows(min_row=2, min_col=2):
            for cell in row:
                if isinstance(cell.value, (int, float)):
                    cell.value = f"{cell.value*100:.2f}%"

        # Create table
        table = Table(displayName="SimilarityTable", 
                     ref=f"A1:{chr(65 + len(doc_names))}{len(doc_names) + 1}")
        table.tableStyleInfo = TableStyleInfo(name="TableStyleMedium9", 
                                           showRowStripes=True)
        sheet.add_table(table)

        # Adjust columns
        for col in sheet.columns:
            max_len = max(len(str(cell.value)) for cell in col)
            sheet.column_dimensions[col[0].column_letter].width = max_len + 2

    return report_filename

def list_reports():
    reports = [f for f in os.listdir(current_app.config['REPORT_FOLDER']) 
              if f.startswith('similarity_report_')]
    return sorted(reports, 
                key=lambda x: os.path.getmtime(
                    os.path.join(current_app.config['REPORT_FOLDER'], x)),
                reverse=True)

def clear_uploads():
    for filename in os.listdir(current_app.config['UPLOAD_FOLDER']):
        file_path = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
        if os.path.isfile(file_path):
            os.remove(file_path)

# Blueprint routes
@web_bp.route('/page5')
def index():
    return render_template('rationalisation.html',
        uploaded_files=sorted(
            f for f in os.listdir(current_app.config['UPLOAD_FOLDER']) 
            if f.endswith('.docx')),
        reports=list_reports(),
        api_key=session.get('api_key', ''),
        model=session.get('model', 'embed-english-light-v2.0'))

@web_bp.route('/upload', methods=['POST'])
def upload_files():
    if 'files' not in request.files:
        flash('No files selected', 'error')
        return redirect(url_for('web.index'))

    files = request.files.getlist('files')
    if not files or files[0].filename == '':
        flash('No files selected', 'error')
        return redirect(url_for('web.index'))

    upload_count = 0
    for file in files:
        if file and allowed_file(file.filename):
            filename = secure_filename(file.filename)
            file.save(os.path.join(current_app.config['UPLOAD_FOLDER'], filename))
            upload_count += 1

    flash(f'{upload_count} files uploaded successfully', 'success')
    return redirect(url_for('web.index'))

@web_bp.route('/analyze', methods=['POST'])
def analyze():
    api_key = request.form.get('api_key')
    model = request.form.get('model')
    
    if not api_key:
        flash('API key required', 'error')
        return redirect(url_for('web.index'))
    
    session.update({
        'api_key': api_key,
        'model': model
    })
    
    docs, doc_names = process_uploaded_documents()
    if not docs:
        flash('No documents found', 'error')
        return redirect(url_for('web.index'))
    
    try:
        co = cohere.Client(api_key)
        embeddings = np.array(co.embed(
            texts=docs,
            model=model
        ).embeddings)
        
        similarity = cosine_similarity(embeddings)
        report = save_similarity_report(similarity, doc_names)
        flash(f'Report generated: {report}', 'success')
    except Exception as e:
        flash(f'Analysis failed: {str(e)}', 'error')
    
    return redirect(url_for('web.index'))

@web_bp.route('/download/<filename>')
def download_report(filename):
    return send_from_directory(
        current_app.config['REPORT_FOLDER'],
        filename,
        as_attachment=True
    )

@web_bp.route('/clear-uploads', methods=['POST'])
def clear_uploads_route():
    clear_uploads()
    flash('Uploads cleared', 'success')
    return redirect(url_for('web.index'))

@web_bp.route('/set-config', methods=['POST'], endpoint='set_config')
def handle_set_config():  # Renamed for clarity
    api_key = request.form.get('api_key')
    model = request.form.get('model')
    
    session['api_key'] = api_key
    session['model'] = model
    
    flash('Configuration saved', 'success')
    return redirect(url_for('web.index'))